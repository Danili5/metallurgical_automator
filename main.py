import os
import sys
import time
import logging
import pandas as pd
from pathlib import Path
from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler

logging.basicConfig(
    level=logging.INFO,
    format="[Time] %(asctime)s\n[Message] %(message)s\n",
    datefmt="%H:%M:%S"
)

class Main(FileSystemEventHandler):
    def on_created(self, event): 
        time.sleep(0.5)
        logging.info(f"The following was created:\n> {event.src_path}")
        
        for report in Path(REPORTS_DIRECTORY).rglob("*.docx"):
            if Path(report).stem.lower().replace(" ", "") == Path(event.src_path).stem.lower().replace(" ", ""):
                try:
                    template = Document(report)

                    data = {f"+{key} {sub_key}+": sub_value for key, value in pd.read_csv(Path(event.src_path), skiprows = [0,2], index_col = 0).to_dict().items() for sub_key, sub_value in value.items()}

                    for row in template.tables[1].rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:        
                                for key, value in data.items():
                                    if key in paragraph.text:
                                        try:
                                            if "." not in value:
                                                value = f"{int(value):,}"
                                        except Exception:
                                            print(Exception)

                                        new_text = paragraph.text.replace(key, str(value))
                                        paragraph.text = ""
                                        paragraph.text = new_text

                                        for run in paragraph.runs:
                                            run.font.name = "Bookman Old Style"
                                            run.font.size = Pt(10)

                                if "+" in paragraph.text:
                                    paragraph.text = ""

                    template.save(report)
                except PermissionError:
                    logging.info("permission error, still running the program observation")

if __name__ == "__main__":    
    load_dotenv()

    DATA_DIRECTORY = os.getenv("DATA_DIRECTORY")
    REPORTS_DIRECTORY = os.getenv("REPORTS_DIRECTORY")

    if not DATA_DIRECTORY or not REPORTS_DIRECTORY:
        logging.error("A directory is missing in .env")
        sys.exit(1)
    
    logging.info(f"Started observation on the following directories:\n> {DATA_DIRECTORY}\n> {REPORTS_DIRECTORY}")

    observer = Observer()

    observer.schedule(Main(), DATA_DIRECTORY, recursive = True)
    observer.start()

    try:
        while observer.is_alive():
            observer.join(1)
    finally:
        observer.stop()
        observer.join()